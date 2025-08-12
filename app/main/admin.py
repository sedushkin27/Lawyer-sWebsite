from django.contrib import admin
from main.models import Review, LegalDocument, AboutMe, AdvantagesWorkingWithMe
from docx import Document
from django.core.files.base import ContentFile
from io import BytesIO
from ckeditor.widgets import CKEditorWidget
from django import forms

class LegalDocumentForm(forms.ModelForm):
    class Meta:
        model = LegalDocument
        fields = ['document_type', 'content', 'file']
        widgets = {
            'content': CKEditorWidget(),
        }

    def clean_file(self):
        file = self.cleaned_data.get('file')
        if file and file.name:
            if not file.name.endswith(('.docx', '.txt')):
                raise forms.ValidationError("Файл должен быть в формате .docx или .txt")
        return file

    def save(self, commit=True):
        instance = super().save(commit=False)
        file = self.cleaned_data.get('file')
        if file and file.name.endswith('.docx'):
            doc = Document(file)
            content = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                if para.style.name.startswith('Heading'):
                    content.append(f'<h{para.style.name[-1]}>{text}</h{para.style.name[-1]}>')
                elif para.runs and any(run.bold for run in para.runs):
                    content.append(f'<strong>{text}</strong>')
                elif para.runs and any(run.italic for run in para.runs):
                    content.append(f'<em>{text}</em>')
                else:
                    content.append(f'<p>{text}</p>')
            instance.content = '\n'.join(content)
        elif file and file.name.endswith('.txt'):
            instance.content = file.read().decode('utf-8').replace('\n', '<br>')
        if commit:
            instance.save()
        return instance
    
class AboutMeForm(forms.ModelForm):
    class Meta:
        model = AboutMe
        fields = '__all__'
        widgets = {
            'main_about_me_text': CKEditorWidget(),
            'other_about_me_text': CKEditorWidget(),
        }

class AdvantagesWorkingWithMeInline(admin.TabularInline):
    model = AdvantagesWorkingWithMe
    extra = 0
    fields = ['title', 'text']


@admin.register(Review)
class ReviewAdmin(admin.ModelAdmin):
    list_display = ['name', 'stars', 'text']
    search_fields = ['name', 'text']

@admin.register(LegalDocument)
class LegalDocumentAdmin(admin.ModelAdmin):
    form = LegalDocumentForm
    list_display = ['document_type', 'updated_at']
    search_fields = ['document_type', 'content']

    def download_as_docx(self, request, queryset):
        for obj in queryset:
            doc = Document()
            doc.add_heading(obj.title, 0)
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(obj.content, 'html.parser')
            for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'strong', 'em']):
                if element.name.startswith('h'):
                    doc.add_heading(element.get_text(), int(element.name[-1]))
                elif element.name == 'strong':
                    para = doc.add_paragraph()
                    run = para.add_run(element.get_text())
                    run.bold = True
                elif element.name == 'em':
                    para = doc.add_paragraph()
                    run = para.add_run(element.get_text())
                    run.italic = True
                else:
                    doc.add_paragraph(element.get_text())
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            response = HttpResponse(
                content=buffer.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename={obj.title}.docx'
            return response
    download_as_docx.short_description = "Скачать как .docx"
    actions = ['download_as_docx']

@admin.register(AboutMe)
class AboutMeAdmin(admin.ModelAdmin):
    form = AboutMeForm
    list_display = ['name', 'sername', 'years_of_experience', 'winning_cases_count', 'client_count']
    search_fields = ['name', 'sername']
    inlines = [AdvantagesWorkingWithMeInline]

    def get_form(self, request, obj=None, **kwargs):
        form = super().get_form(request, obj, **kwargs)
        form.base_fields['first_photo'].widget.attrs.update({'accept': 'image/*'})
        form.base_fields['second_photo'].widget.attrs.update({'accept': 'image/*'})
        form.base_fields['third_photo'].widget.attrs.update({'accept': 'image/*'})
        return form